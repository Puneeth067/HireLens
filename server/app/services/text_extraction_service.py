"""
Text Extraction Service
Handles extraction of text from PDF and DOCX files
"""
import os
import logging
from typing import Optional
import PyPDF2
import fitz  # PyMuPDF
from docx import Document
import mammoth
from pathlib import Path

logger = logging.getLogger(__name__)

class TextExtractionService:
    """Service for extracting text from various document formats"""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc'}
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a document file
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
            Exception: If extraction fails
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using multiple methods"""
        text = ""
        
        # Method 1: Try PyMuPDF (fitz) first - better for complex PDFs
        try:
            text = self._extract_pdf_with_fitz(file_path)
            if text.strip():
                logger.info(f"Successfully extracted PDF text using PyMuPDF: {file_path}")
                return text
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed for {file_path}: {str(e)}")
        
        # Method 2: Fallback to PyPDF2
        try:
            text = self._extract_pdf_with_pypdf2(file_path)
            if text.strip():
                logger.info(f"Successfully extracted PDF text using PyPDF2: {file_path}")
                return text
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed for {file_path}: {str(e)}")
        
        # If both methods fail but no exception was raised, return empty string
        logger.error(f"All PDF extraction methods failed for {file_path}")
        return ""
    
    def _extract_pdf_with_fitz(self, file_path: str) -> str:
        """Extract text using PyMuPDF (fitz)"""
        text = ""
        doc = fitz.open(file_path)
        
        try:
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text()
                text += "\n"  # Add page break
            
        finally:
            doc.close()
        
        return self._clean_text(text)
    
    def _extract_pdf_with_pypdf2(self, file_path: str) -> str:
        """Extract text using PyPDF2"""
        text = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
                text += "\n"  # Add page break
        
        return self._clean_text(text)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX/DOC file using multiple methods"""
        
        # Method 1: Try python-docx first
        try:
            text = self._extract_docx_with_python_docx(file_path)
            if text.strip():
                logger.info(f"Successfully extracted DOCX text using python-docx: {file_path}")
                return text
        except Exception as e:
            logger.warning(f"python-docx extraction failed for {file_path}: {str(e)}")
        
        # Method 2: Fallback to mammoth (better for complex formatting)
        try:
            text = self._extract_docx_with_mammoth(file_path)
            if text.strip():
                logger.info(f"Successfully extracted DOCX text using mammoth: {file_path}")
                return text
        except Exception as e:
            logger.warning(f"Mammoth extraction failed for {file_path}: {str(e)}")
        
        logger.error(f"All DOCX extraction methods failed for {file_path}")
        return ""
    
    def _extract_docx_with_python_docx(self, file_path: str) -> str:
        """Extract text using python-docx"""
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return self._clean_text(text)
    
    def _extract_docx_with_mammoth(self, file_path: str) -> str:
        """Extract text using mammoth"""
        with open(file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            text = result.value
            
            # Log any warnings from mammoth
            if result.messages:
                for message in result.messages:
                    logger.warning(f"Mammoth warning: {message}")
        
        return self._clean_text(text)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        import re
        
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Replace multiple newlines with maximum of 2
        text = re.sub(r'\n\n+', '\n\n', text)
        
        # Remove trailing/leading whitespace from each line
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines]
        text = '\n'.join(cleaned_lines)
        
        # Remove excessive whitespace at start and end
        text = text.strip()
        
        return text
    
    def get_file_info(self, file_path: str) -> dict:
        """Get information about the file"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_stats = os.stat(file_path)
        file_extension = Path(file_path).suffix.lower()
        
        info = {
            'filename': os.path.basename(file_path),
            'file_size': file_stats.st_size,
            'file_extension': file_extension,
            'is_supported': file_extension in self.supported_formats,
            'last_modified': file_stats.st_mtime
        }
        
        # Add format-specific information
        if file_extension == '.pdf':
            info.update(self._get_pdf_info(file_path))
        elif file_extension in ['.docx', '.doc']:
            info.update(self._get_docx_info(file_path))
        
        return info
    
    def _get_pdf_info(self, file_path: str) -> dict:
        """Get PDF-specific information"""
        try:
            doc = fitz.open(file_path)
            info = {
                'page_count': doc.page_count,
                'has_text': False
            }
            
            # Check if PDF has extractable text
            for page_num in range(min(3, doc.page_count)):  # Check first 3 pages
                page = doc[page_num]
                text = page.get_text().strip()
                if text:
                    info['has_text'] = True
                    break
            
            doc.close()
            return info
            
        except Exception as e:
            logger.warning(f"Could not get PDF info for {file_path}: {str(e)}")
            return {'page_count': 0, 'has_text': False}
    
    def _get_docx_info(self, file_path: str) -> dict:
        """Get DOCX-specific information"""
        try:
            doc = Document(file_path)
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            return {
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'has_content': paragraph_count > 0 or table_count > 0
            }
            
        except Exception as e:
            logger.warning(f"Could not get DOCX info for {file_path}: {str(e)}")
            return {'paragraph_count': 0, 'table_count': 0, 'has_content': False}